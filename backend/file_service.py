import os
import uuid
import re
from pathlib import Path

_BASE_CANDIDATE = "/data/workspace"
try:
    os.makedirs(_BASE_CANDIDATE, exist_ok=True)
    BASE = _BASE_CANDIDATE
except (PermissionError, OSError):
    BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'workspace')
    os.makedirs(BASE, exist_ok=True)

TEXT_EXTS = {'txt', 'md', 'py', 'js', 'ts', 'jsx', 'tsx', 'json', 'csv', 'html', 'css', 'xml', 'yaml', 'yml', 'ini', 'cfg', 'log', 'sh', 'bat', 'ps1', 'env', 'rst', 'tex', 'c', 'cpp', 'h', 'java', 'rs', 'go', 'rb', 'php', 'swift', 'kt', 'scala', 'sql', 'r', 'lua'}
HEADING_PATTERN = re.compile(r'^(#{1,3})\s+\w')
BULLET_PATTERN = re.compile(r'^\s*[-*]\s+\S')


def get_workspace(session_id: str) -> str:
    ws = os.path.join(BASE, session_id)
    os.makedirs(ws, exist_ok=True)
    return ws


def _safe_path(filename: str, session_id: str = "default") -> Path:
    ws = get_workspace(session_id)
    abs_path = os.path.abspath(os.path.join(ws, filename))
    if not abs_path.startswith(os.path.abspath(ws)):
        raise ValueError("Path traversal detected")
    return Path(abs_path)


def _extract_text_from_pdf(path: Path) -> str:
    try:
        import PyPDF2
        text = ""
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except ImportError:
        return "[PDF text extraction requires PyPDF2]"
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except ImportError:
        return "[DOCX text extraction requires python-docx]"
    except Exception as e:
        return f"[DOCX extraction error: {e}]"


def extract_text(file_path: str, session_id: str = "default") -> str | None:
    path = _safe_path(file_path, session_id)
    if not path.exists():
        return None
    ext = path.suffix.lower().lstrip('.')
    if ext == 'pdf':
        return _extract_text_from_pdf(path)
    if ext == 'docx':
        return _extract_text_from_docx(path)
    if ext in TEXT_EXTS:
        try:
            return path.read_text(encoding='utf-8').strip()
        except Exception:
            return None
    return None


def save_upload(file_bytes: bytes, original_name: str, session_id: str = "default") -> dict:
    safe_name = os.path.basename(original_name)
    if not safe_name:
        safe_name = f"unnamed_{uuid.uuid4()[:8]}"
    path = _safe_path(safe_name, session_id)
    path.write_bytes(file_bytes)
    size = len(file_bytes)
    text = extract_text(safe_name, session_id)
    result = {
        "file_id": safe_name,
        "name": safe_name,
        "saved_as": safe_name,
        "size": size,
        "path": str(path),
    }
    if text:
        result["extracted_text"] = text
    return result


def get_file_path(filename: str, session_id: str = "default") -> Path | None:
    path = _safe_path(filename, session_id)
    if path.exists():
        return path
    return None


def read_file_content(filename: str, session_id: str = "default") -> str:
    path = _safe_path(filename, session_id)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {filename}")
    return path.read_text(encoding='utf-8')


def write_file_content(filename: str, content: str, session_id: str = "default") -> str:
    path = _safe_path(filename, session_id)
    path.write_text(content, encoding='utf-8')
    return str(path)


def list_files(session_id: str = "default") -> list[dict]:
    ws = get_workspace(session_id)
    files = []
    for f in sorted(os.listdir(ws)):
        fp = os.path.join(ws, f)
        if os.path.isfile(fp):
            files.append({
                "name": f,
                "size": os.path.getsize(fp),
                "modified": os.path.getmtime(fp),
            })
    return files


def delete_file(filename: str, session_id: str = "default") -> bool:
    try:
        path = _safe_path(filename, session_id)
        if path.exists():
            path.unlink()
            return True
    except ValueError:
        pass
    return False


def create_pdf(content: str, name: str = "report.pdf", session_id: str = "default") -> str:
    try:
        from fpdf import FPDF
        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_font("DejaVu", "", "DejaVuSansCondensed.ttf", uni=True)
        pdf.set_font("DejaVu", size=11)
        pdf.multi_cell(0, 8, "Report", align="C")
        pdf.ln(4)
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                pdf.ln(2)
                continue
            try:
                pdf.multi_cell(0, 8, line)
            except RuntimeError:
                pdf.ln(2)
                continue
        path = os.path.join(get_workspace(session_id), name)
        pdf.output(path)
        return name
    except ImportError:
        raise RuntimeError("fpdf2 not installed")


def create_docx(content: str, name: str = "report.docx", session_id: str = "default") -> str:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Report', 0)
        skip_md = name.lower().endswith(('.py', '.js', '.ts', '.sh', '.html', '.css', '.json'))
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if skip_md:
                doc.add_paragraph(line)
            elif HEADING_PATTERN.match(line):
                hashes = len(line.split()[0])
                doc.add_heading(line.lstrip('# '), hashes)
            elif BULLET_PATTERN.match(line):
                doc.add_paragraph(line.lstrip('- ').lstrip('* '), style='List Bullet')
            else:
                doc.add_paragraph(line)
        path = os.path.join(get_workspace(session_id), name)
        doc.save(path)
        return name
    except ImportError:
        raise RuntimeError("python-docx not installed")


def create_txt(content: str, name: str = "report.txt", session_id: str = "default") -> str:
    ws = get_workspace(session_id)
    path = os.path.join(ws, name)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)
    return name
